from tkinter import *
from docx import Document
from tkinter import messagebox

def save_resume():
    name = name_entry.get()
    age = age_entry.get()
    dob = dob_entry.get()
    gender = gender_var.get()
    qualification = qual_entry.get()
    address = addr_entry.get()

    # Create a Word document
    doc = Document()
    doc.add_heading("Resume", level=1)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"DOB: {dob}")
    doc.add_paragraph(f"Gender: {gender}")
    doc.add_paragraph(f"Qualification: {qualification}")
    doc.add_paragraph(f"Address: {address}")
    doc.save("resume.docx")

    messagebox.showinfo("Success", "Resume saved as resume.docx")

# GUI Window
root = Tk()
root.title("Resume Creator")

Label(root, text="Name:").grid(row=0, column=0, pady=5, sticky=W)
Label(root, text="Age:").grid(row=1, column=0, pady=5, sticky=W)
Label(root, text="DOB:").grid(row=2, column=0, pady=5, sticky=W)
Label(root, text="Gender:").grid(row=3, column=0, pady=5, sticky=W)
Label(root, text="Qualification:").grid(row=4, column=0, pady=5, sticky=W)
Label(root, text="Address:").grid(row=5, column=0, pady=5, sticky=W)

name_entry = Entry(root)
age_entry = Entry(root)
dob_entry = Entry(root)
qual_entry = Entry(root)
addr_entry = Entry(root)

name_entry.grid(row=0, column=1)
age_entry.grid(row=1, column=1)
dob_entry.grid(row=2, column=1)
qual_entry.grid(row=4, column=1)
addr_entry.grid(row=5, column=1)

# Gender radio buttons
gender_var = StringVar()
gender_var.set("Male")

Radiobutton(root, text="Male", variable=gender_var, value="Male").grid(row=3, column=1, sticky=W)
Radiobutton(root, text="Female", variable=gender_var, value="Female").grid(row=3, column=2, sticky=W)

# Save Button
Button(root, text="Save Resume", command=save_resume).grid(row=6, column=0, columnspan=3, pady=10)

root.mainloop()
